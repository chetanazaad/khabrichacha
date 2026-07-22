"""
KhabriChacha — Report Exporter

Creates report.md, report.txt, report.json, a professional A4 report.pdf,
and a report.docx (Word). Uses reportlab for PDF generation and
python-docx for Word generation.

This module ONLY creates files. It never manages project folders.
ProjectManager is responsible for storing the output.
"""

import json
from datetime import datetime
from io import BytesIO
from typing import Any, Dict, List, Optional
from xml.sax.saxutils import escape as _xml_escape
from loguru import logger

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm, mm
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        PageBreak, HRFlowable,
    )
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    logger.warning("reportlab not installed — PDF generation will be skipped.")

try:
    import docx
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not installed — Word (.docx) generation will be skipped.")


def _esc(text: Any) -> str:
    """
    Escape text before it's interpolated into a ReportLab Paragraph, which
    parses its input as a small XML-like markup language. Without this, any
    finding/evidence/source text containing a literal '<', '>', or '&' —
    extremely common in scraped web content — raises a parse error and
    silently kills PDF generation (the caller's try/except just swallows
    it, so the user simply never gets a PDF). Only call this on *data*, not
    on the literal "<b>"/"<br/>" tags this module writes itself.
    """
    return _xml_escape(str(text) if text is not None else "")


class ReportExporter:
    """
    Generates research deliverables from orchestrator results.

    Returns a dict with keys:
        "report_md": str
        "report_json": dict
        "report_pdf_bytes": bytes | None
        "report_docx_bytes": bytes | None
    """

    def generate(
        self,
        *,
        title: str,
        mission: str,
        provider: str = "",
        model: str = "",
        findings: List[str],
        sources: List[Dict[str, str]],
        evidence: str = "",
        research_state: Optional[Dict[str, Any]] = None,
        timeline: str = "",
    ) -> Dict[str, Any]:
        """Produce all report formats (md, json, pdf, docx)."""

        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        # ── Markdown ─────────────────────────────────────────
        report_md = self._build_markdown(
            title=title, mission=mission, provider=provider,
            model=model, findings=findings, sources=sources,
            evidence=evidence, timeline=timeline, timestamp=timestamp,
        )

        # ── JSON ─────────────────────────────────────────────
        report_json = {
            "title": title,
            "mission": mission,
            "provider": provider,
            "model": model,
            "generated": timestamp,
            "findings": findings,
            "sources": sources,
            "evidence_summary": evidence[:500] if evidence else "",
            "timeline": timeline,
            "research_state": research_state or {},
        }

        # ── PDF ──────────────────────────────────────────────
        report_pdf_bytes = None
        if REPORTLAB_AVAILABLE:
            try:
                report_pdf_bytes = self._build_pdf(
                    title=title, mission=mission, provider=provider,
                    model=model, findings=findings, sources=sources,
                    evidence=evidence, timeline=timeline, timestamp=timestamp,
                )
            except Exception as e:
                logger.error(f"PDF generation failed: {e}")
        else:
            logger.warning("Skipping PDF generation — reportlab not available.")

        # ── DOCX (Word) ────────────────────────────────────────
        report_docx_bytes = None
        if DOCX_AVAILABLE:
            try:
                report_docx_bytes = self._build_docx(
                    title=title, mission=mission, provider=provider,
                    model=model, findings=findings, sources=sources,
                    evidence=evidence, timeline=timeline, timestamp=timestamp,
                )
            except Exception as e:
                logger.error(f"DOCX generation failed: {e}")
        else:
            logger.warning("Skipping DOCX generation — python-docx not available.")

        return {
            "report_md": report_md,
            "report_json": report_json,
            "report_pdf_bytes": report_pdf_bytes,
            "report_docx_bytes": report_docx_bytes,
        }

    # ── Markdown builder ─────────────────────────────────────

    @staticmethod
    def _build_markdown(
        *, title, mission, provider, model, findings,
        sources, evidence, timeline, timestamp,
    ) -> str:
        lines = [
            f"# {title}",
            "",
            "## Executive Summary",
        ]
        summary_points = [s.strip().rstrip(".") + "." for s in findings[:3] if s.strip()]
        if summary_points:
            lines.append(
                "This report consolidates key research findings. "
                "Highlights include: " + " ".join(summary_points)
            )
        elif evidence:
            lines.append(evidence[:400] + ("..." if len(evidence) > 400 else ""))
        else:
            lines.append("No conclusive findings were recorded.")
        lines.append("")

        lines += ["## Research Mission", mission or "(not specified)", ""]
        lines += ["## Methodology",
                  f"- **Provider**: {provider or 'N/A'}",
                  f"- **Model**: {model or 'N/A'}",
                  f"- **Generated**: {timestamp}", ""]

        if timeline:
            lines += ["## Research Timeline", timeline, ""]

        lines.append("## Key Findings")
        if findings:
            for i, f in enumerate(findings, 1):
                lines.append(f"{i}. {f}")
        elif evidence:
            lines.append(evidence[:1500])
        else:
            lines.append("No findings provided.")
        lines.append("")

        if evidence:
            lines += ["## Evidence Summary", evidence, ""]

        lines.append("## Sources")
        if sources:
            for s in sources:
                t = s.get("title", "Untitled")
                u = s.get("url", "#")
                lines.append(f"- [{t}]({u})")
        else:
            lines.append("No sources provided.")
        lines.append("")

        lines += [
            "## Research Statistics",
            f"- **Total Findings**: {len(findings)}",
            f"- **Total Sources**: {len(sources)}",
            f"- **Generated**: {timestamp}",
        ]
        return "\n".join(lines)

    # ── PDF builder ──────────────────────────────────────────

    @staticmethod
    def _build_pdf(
        *, title, mission, provider, model, findings,
        sources, evidence, timeline, timestamp,
    ) -> bytes:
        buf = BytesIO()
        doc = SimpleDocTemplate(
            buf, pagesize=A4,
            leftMargin=2 * cm, rightMargin=2 * cm,
            topMargin=2 * cm, bottomMargin=2 * cm,
        )

        styles = getSampleStyleSheet()
        # Custom styles
        styles.add(ParagraphStyle(
            "CoverTitle", parent=styles["Title"],
            fontSize=28, leading=34,
            textColor=HexColor("#1a1a2e"),
            alignment=TA_CENTER, spaceAfter=12,
        ))
        styles.add(ParagraphStyle(
            "CoverSub", parent=styles["Normal"],
            fontSize=14, leading=18,
            textColor=HexColor("#4a4a6a"),
            alignment=TA_CENTER, spaceAfter=6,
        ))
        styles.add(ParagraphStyle(
            "SectionHead", parent=styles["Heading2"],
            fontSize=16, leading=20,
            textColor=HexColor("#16213e"),
            spaceBefore=18, spaceAfter=8,
        ))
        styles.add(ParagraphStyle(
            "BodyText2", parent=styles["BodyText"],
            fontSize=10, leading=14,
            alignment=TA_JUSTIFY,
        ))
        styles.add(ParagraphStyle(
            "BulletItem", parent=styles["BodyText"],
            fontSize=10, leading=14, leftIndent=18,
            bulletIndent=6,
        ))

        story = []

        # ── Cover Page ───────────────────────────────────────
        story.append(Spacer(1, 6 * cm))
        story.append(Paragraph(_esc(title), styles["CoverTitle"]))
        story.append(Spacer(1, 0.5 * cm))
        story.append(Paragraph(f"Research Mission: {_esc(mission or 'N/A')}", styles["CoverSub"]))
        story.append(Paragraph(f"Generated: {_esc(timestamp)}", styles["CoverSub"]))
        story.append(Paragraph(f"Provider: {_esc(provider or 'N/A')}  |  Model: {_esc(model or 'N/A')}", styles["CoverSub"]))
        story.append(PageBreak())

        # Helper to clean markdown headers and HTML entities for ReportLab PDF paragraphs
        def _clean_pdf_text(raw_text: str) -> str:
            if not raw_text:
                return ""
            clean = raw_text.replace("&nbsp;", " ").replace("### ", "").replace("## ", "").replace("# ", "")
            return _esc(clean)

        # ── Executive Summary ────────────────────────────────
        story.append(Paragraph("Executive Summary", styles["SectionHead"]))
        summary_points = [s.strip().rstrip(".") + "." for s in findings[:3] if s.strip()]
        if summary_points:
            summary_text = (
                "This report consolidates key research findings. "
                "Highlights include: " + " ".join(summary_points)
            )
        elif evidence:
            summary_text = evidence[:400] + ("..." if len(evidence) > 400 else "")
        else:
            summary_text = "No conclusive findings were recorded."
        story.append(Paragraph(_clean_pdf_text(summary_text), styles["BodyText2"]))
        story.append(Spacer(1, 0.5 * cm))

        # ── Methodology ──────────────────────────────────────
        story.append(Paragraph("Methodology", styles["SectionHead"]))
        story.append(Paragraph(f"<b>Provider:</b> {_esc(provider or 'N/A')}", styles["BodyText2"]))
        story.append(Paragraph(f"<b>Model:</b> {_esc(model or 'N/A')}", styles["BodyText2"]))
        story.append(Paragraph(f"<b>Generated:</b> {_esc(timestamp)}", styles["BodyText2"]))
        story.append(Spacer(1, 0.3 * cm))

        # ── Research Timeline ────────────────────────────────
        if timeline:
            story.append(Paragraph("Research Timeline", styles["SectionHead"]))
            for line in timeline.split("\n"):
                if line.strip():
                    story.append(Paragraph(_clean_pdf_text(line.strip()), styles["BodyText2"]))
            story.append(Spacer(1, 0.3 * cm))

        # ── Key Findings ─────────────────────────────────────
        story.append(Paragraph("Key Findings", styles["SectionHead"]))
        if findings:
            for i, f in enumerate(findings, 1):
                text = f[:500] + "..." if len(f) > 500 else f
                story.append(Paragraph(f"{i}. {_clean_pdf_text(text)}", styles["BulletItem"]))
        elif evidence:
            story.append(Paragraph(_clean_pdf_text(evidence[:1500]), styles["BodyText2"]))
        else:
            story.append(Paragraph("No findings provided.", styles["BodyText2"]))
        story.append(Spacer(1, 0.3 * cm))

        # ── Evidence Summary ─────────────────────────────────
        if evidence:
            story.append(Paragraph("Evidence Summary", styles["SectionHead"]))
            # Keep evidence reasonable for the PDF
            ev_text = evidence[:2000] + "..." if len(evidence) > 2000 else evidence
            for line in ev_text.split("\n"):
                if line.strip():
                    story.append(Paragraph(_esc(line.strip()), styles["BodyText2"]))
            story.append(Spacer(1, 0.3 * cm))

        # ── Evidence Statistics ──────────────────────────────
        story.append(Paragraph("Evidence Statistics", styles["SectionHead"]))
        stats_data = [
            ["Metric", "Value"],
            ["Total Findings", str(len(findings))],
            ["Total Sources", str(len(sources))],
            ["Generated", timestamp],
        ]
        stats_table = Table(stats_data, colWidths=[8 * cm, 8 * cm])
        stats_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), HexColor("#16213e")),
            ("TEXTCOLOR", (0, 0), (-1, 0), HexColor("#ffffff")),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 10),
            ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#cccccc")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [HexColor("#f8f8f8"), HexColor("#ffffff")]),
            ("ALIGN", (0, 0), (-1, -1), "LEFT"),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        story.append(stats_table)
        story.append(Spacer(1, 0.5 * cm))

        # ── References ───────────────────────────────────────
        story.append(Paragraph("References", styles["SectionHead"]))
        if sources:
            for s in sources:
                t = s.get("title", "Untitled")
                u = s.get("url", "")
                ref_text = f"<b>{_esc(t)}</b>"
                if u:
                    ref_text += f"<br/><font size='8' color='#4a4a6a'>{_esc(u)}</font>"
                story.append(Paragraph(ref_text, styles["BulletItem"]))
        else:
            story.append(Paragraph("No references collected.", styles["BodyText2"]))

        # ── Build ────────────────────────────────────────────
        doc.build(story)
        return buf.getvalue()

    # ── DOCX (Word) builder ───────────────────────────────────

    @staticmethod
    def _build_docx(
        *, title, mission, provider, model, findings,
        sources, evidence, timeline, timestamp,
    ) -> bytes:
        document = docx.Document()

        # Title
        h = document.add_heading(title or "Research Report", level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        sub = document.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub.add_run(f"Generated: {timestamp}  |  Provider: {provider or 'N/A'}  |  Model: {model or 'N/A'}")
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x4a, 0x4a, 0x6a)

        document.add_heading("Executive Summary", level=1)
        summary_points = [s.strip().rstrip(".") + "." for s in findings[:3] if s.strip()]
        if summary_points:
            document.add_paragraph(
                "This report consolidates key research findings. "
                "Highlights include: " + " ".join(summary_points)
            )
        else:
            document.add_paragraph("No conclusive findings were recorded.")

        document.add_heading("Research Mission", level=1)
        document.add_paragraph(mission or "(not specified)")

        document.add_heading("Methodology", level=1)
        document.add_paragraph(f"Provider: {provider or 'N/A'}")
        document.add_paragraph(f"Model: {model or 'N/A'}")
        document.add_paragraph(f"Generated: {timestamp}")

        if timeline:
            document.add_heading("Research Timeline", level=1)
            for line in timeline.split("\n"):
                if line.strip():
                    document.add_paragraph(line.strip())

        document.add_heading("Key Findings", level=1)
        if findings:
            for f in findings:
                document.add_paragraph(f, style="List Number")
        else:
            document.add_paragraph("No findings provided.")

        if evidence:
            document.add_heading("Evidence Summary", level=1)
            ev_text = evidence[:4000] + "..." if len(evidence) > 4000 else evidence
            for line in ev_text.split("\n"):
                if line.strip():
                    document.add_paragraph(line.strip())

        document.add_heading("Evidence Statistics", level=1)
        table = document.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text = "Metric", "Value"
        for metric, value in [
            ("Total Findings", str(len(findings))),
            ("Total Sources", str(len(sources))),
            ("Generated", timestamp),
        ]:
            row = table.add_row().cells
            row[0].text, row[1].text = metric, value

        document.add_heading("References", level=1)
        if sources:
            for s in sources:
                t = s.get("title", "Untitled")
                u = s.get("url", "")
                p = document.add_paragraph(style="List Bullet")
                run = p.add_run(t)
                run.bold = True
                if u:
                    p.add_run(f"\n{u}").italic = True
        else:
            document.add_paragraph("No references collected.")

        buf = BytesIO()
        document.save(buf)
        return buf.getvalue()
